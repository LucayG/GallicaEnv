# -*- coding: utf-8 -*-

import os
import ast
try :
	import pandas as pd
except ModuleNotFoundError :
	os.system("pip install pandas")
	import pandas as pd
try :
	from docx import Document
	from docx.text.paragraph import Paragraph
except ModuleNotFoundError :
	os.system("pip install python-docx")
	from docx import Document
	from docx.text.paragraph import Paragraph

nom_file_xml = input("Entrez le nom du fichier xml (avec l'extension .xml) : ")
nom_file_extraits = input("Entrez le nom du fichier txt contenant les extraits (avec l'extension .txt) : ")

df = pd.read_csv(nom_file_xml, encoding = 'utf8')

dico_infos= {}
def recuperation_infos(data):
    page = data["Page"]
    lien_ark = data["ARK"]
    occ_mots = ast.literal_eval(data["Occurance Mot"])
    titre = data["Titre"]
    auteur = data["Auteur"]
    date = data["Date d'édition"]
    lien_page = data["Lien Gallica Page"]
    dico_infos[data["Match"].split('-')[1]] = {"page" : page, "ark" : lien_ark, "occ" : occ_mots, "titre" : titre, "auteur" : auteur, "date" : date, "lien_page" : lien_page}
   
df.apply(recuperation_infos, axis=1)

dico_extraits = {}
with open(nom_file_extraits, 'r', encoding = 'utf8') as fic:
	for line in fic.readlines():
		line = line.split("-")
		l = line[0].split('_')[0]
		p = line[0].split('_')[1]
		d = line[0].split('_')[2]
		p = l[0:3]+"E_"+p
		for key in dico_infos:
			if dico_infos[key]["page"] == p and dico_infos[key]["ark"] == d:
				dico_extraits[key] = '-'.join(line[1:])
		

document = Document()
document.add_heading('Résultats de la requête', 0)
for key in dico_infos:
  p = document.add_paragraph("")
  p = document.add_paragraph("")
  p.add_run("Match "+key).bold = True
  # titre
  p = document.add_paragraph("")
  p.add_run("Titre : ").italic = True
  p.add_run(str(dico_infos[key]["titre"])+"\n")
  # auteur
  p.add_run("Auteur : ").italic = True
  p.add_run(str(dico_infos[key]["auteur"])+"\n")
  # date
  p.add_run("Date d'édition : ").italic = True
  p.add_run(str(dico_infos[key]["date"])+"\n")
  # ark
  p.add_run("Ark du document : ").italic = True
  p.add_run(str(dico_infos[key]["ark"])+"\n")
  # lien de la page
  p.add_run("Lien de la page : ").italic = True
  p.add_run(str(dico_infos[key]["lien_page"])+"\n")
  # occurrences
  p.add_run("Nombre d'occurrences des mots trouvés : \n").italic = True
  for k in dico_infos[key]["occ"]:
    p.add_run(k).bold = True
    p.add_run(" : "+str(dico_infos[key]["occ"][k])+"\n")
  # extrait
  p.add_run("Extrait : ").italic = True
  if key in dico_extraits :
  	p.add_run("\n" + str(dico_extraits[key]))
  # saut de ligne
  p = document.add_paragraph("")

document.save('Resultats_et_extraits.docx')
